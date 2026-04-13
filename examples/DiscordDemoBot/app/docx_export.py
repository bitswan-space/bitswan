"""Generate a filled Demo Registration .docx from current meeting state."""
from __future__ import annotations

import asyncio
import io
from functools import partial

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


def _set_cell_bg(cell, hex_color: str) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _sync_generate(
    meeting_date: str,
    present: list[str],
    absent: list[str],
    all_persons: list[str],
    projects_with_topics: list[tuple[str, str, list[tuple[str, str]]]],
) -> bytes:
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(1.8)
        section.bottom_margin = Cm(1.8)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)

    # Title
    p = doc.add_paragraph()
    run = p.add_run("Demo Meeting — Registration Sheet")
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor.from_string("1F3864")

    doc.add_paragraph(f"Date: {meeting_date}")
    doc.add_paragraph()

    # Attendees
    doc.add_heading("Attendees", level=2)
    att = doc.add_table(rows=1, cols=3)
    att.style = "Table Grid"
    for cell, txt in zip(att.rows[0].cells, ["#", "Name / Email", "Present"]):
        cell.text = txt
        _set_cell_bg(cell, "1F3864")
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.size = Pt(9)

    present_set = set(present)
    for i, name in enumerate(all_persons, 1):
        row = att.add_row()
        row.cells[0].text = str(i)
        row.cells[1].text = name
        row.cells[2].text = "✔" if name in present_set else ""
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(9)
        if name not in present_set:
            _set_cell_bg(row.cells[2], "FFF2CC")

    doc.add_paragraph()

    # Projects & Presentations
    doc.add_heading("Projects & Presentations", level=2)
    proj_tbl = doc.add_table(rows=1, cols=3)
    proj_tbl.style = "Table Grid"
    for cell, txt in zip(proj_tbl.rows[0].cells, ["Project (Guarantor)", "Topic", "Presenter"]):
        cell.text = txt
        _set_cell_bg(cell, "1F3864")
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.size = Pt(9)

    for proj_name, guarantor, topics in projects_with_topics:
        hdr_row = proj_tbl.add_row()
        c = hdr_row.cells[0]
        c.merge(hdr_row.cells[1]).merge(hdr_row.cells[2])
        c.text = f"{proj_name}   —   Guarantor: {guarantor}"
        _set_cell_bg(c, "D9E1F2")
        for para in c.paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.size = Pt(9)

        for title, presenter in topics:
            row = proj_tbl.add_row()
            row.cells[0].text = ""
            row.cells[1].text = title
            row.cells[2].text = presenter
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.size = Pt(9)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


async def generate_docx(
    meeting_date: str,
    present: list[str],
    absent: list[str],
    all_persons: list[str],
    projects_with_topics: list[tuple[str, str, list[tuple[str, str]]]],
) -> bytes:
    loop = asyncio.get_event_loop()
    return await loop.run_in_executor(
        None,
        partial(_sync_generate, meeting_date, present, absent, all_persons, projects_with_topics),
    )
